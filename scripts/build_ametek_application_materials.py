from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import BaseDocTemplate, Frame, KeepTogether, PageTemplate, Paragraph, Spacer


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"

NAME = "Russell Dudek"
PHONE = "412.287.8640"
EMAIL = "russelldudek@gmail.com"
LINKEDIN = "linkedin.com/in/russelldudek"
SITE_URL = "https://russelldudek.github.io/AMETEK/"
CONTACT = f"Pittsburgh, PA | {PHONE} | {EMAIL} | {LINKEDIN} | {SITE_URL}"

NAVY_RGB = RGBColor(23, 50, 77)
ACCENT_RGB = RGBColor(47, 111, 115)
MUTED_RGB = RGBColor(91, 103, 112)
INK_RGB = RGBColor(31, 37, 43)
RULE_RGB = RGBColor(216, 222, 230)

NAVY = colors.HexColor("#17324D")
ACCENT = colors.HexColor("#2F6F73")
MUTED = colors.HexColor("#5B6770")
INK = colors.HexColor("#1F252B")
RULE = colors.HexColor("#D8DEE6")

RESUME = {
    "filename": "Russell_Dudek_AMETEK_PAI_Division_VP_IT_Resume",
    "headline": "Division IT Transformation Leader | AI-First Operations | Manufacturing Systems, Data Governance, and Execution",
    "summary": (
        "AI transformation and operations leader with 20+ years across manufacturing, logistics, robotics, customer operations, "
        "field technology, ERP-adjacent workflows, and executive operating systems. Strongest where decentralized teams need a "
        "common operating spine: clean process ownership, usable data, disciplined governance, practical automation, and human "
        "adoption. Builds the translation layer between business reality and technology execution so transformation becomes "
        "visible, measurable, and durable."
    ),
    "scale": (
        "Selected Scale | 1,000+ deployed machines supported | 5M annual second-day orders | 99.98% on-time delivery | "
        "100% FAA safety record | Led 1,000+ Amazon employees | 20+ years operations"
    ),
    "top_section_title": "AMETEK P&AI Relevance",
    "top_bullets": [
        "Built operating cadence across production, quality, purchasing, support, engineering, training, customer operations, and leadership reporting.",
        "Converted scattered operational knowledge into SOPs, dashboards, escalation paths, onboarding flows, and AI-assisted review loops.",
        "Designed human-in-the-loop AI systems where leaders can see evidence, decision gates, ownership, and operational risk before it becomes noise.",
        "Led high-pressure launch and scale work where customer promise, quality discipline, safety, staffing, and daily metrics had to align.",
        "Brings a business-first IT posture: standardize the core where variance adds cost, preserve local autonomy where it creates real advantage.",
    ],
    "roles": [
        {
            "title": "Director of Operations",
            "company": "Vape-Jet / PDX Tech Services",
            "location": "Pittsburgh, PA / Remote",
            "dates": "2025 - Present",
            "lead": "Lead AI-first operating transformation for a software-enabled hardware company moving through scale, relocation, manufacturing change, and customer-support complexity.",
            "bullets": [
                "Built an AI fabric across operations, GTM, product, software, HR, ERP/Odoo workflows, Jira, fleet intelligence, and leadership reporting.",
                "Prepared West Coast operations for scale while architecting and launching a new East Coast HQ for distributed manufacturing and support.",
                "Advanced proactive support for 1,000+ machines using telemetry, throughput patterns, hidden error signals, and AI-assisted review.",
                "Protected customer trust by turning support pain into clearer engineering, service, production, documentation, and training priorities.",
                "Created purchasing-control and Odoo-based planning infrastructure to make BOM, inventory, incoming, vendor, shortage, and sourcing decisions visible.",
                "Built department hubs, review artifacts, escalation paths, and standard work that reduced dependence on individual memory.",
            ],
        },
        {
            "title": "Founder / AI Transformation Advisor",
            "company": "DudeWorth",
            "location": "Pittsburgh, PA / Remote",
            "dates": "2017 - Present",
            "lead": "Build human plus AI operating systems for strategy, decision support, workflow automation, and governed adoption.",
            "bullets": [
                "Built DudexOS with Core as an agentic harness for research, decision records, QA checks, operating artifacts, and executive workflow support.",
                "Designed AI adoption methods around real work: leadership alignment, policy, training, metrics, human review, and continuous improvement loops.",
                "Created practical AI-supported systems for funnel strategy, knowledge capture, reporting cadence, workflow design, and business development.",
                "Translate AI concepts into operating patterns that improve judgment, visibility, accountability, and trust rather than novelty output.",
            ],
        },
        {
            "title": "Regional Director",
            "company": "Cardinal Building Products",
            "location": "Pittsburgh, PA",
            "dates": "2021 - 2023",
            "lead": "Built and led a remote regional operation across customers, digital systems, logistics, and operating cadence.",
            "bullets": [
                "Strengthened major customer relationships, follow-up cadence, sales coordination, fulfillment support, and regional accountability.",
                "Advanced route planning, CRM, digital marketing, online buying, and regional operating model concepts.",
                "Built remote operating rhythms that kept customer work, internal handoffs, and leadership visibility moving through change.",
            ],
        },
        {
            "title": "Chief Pilot, Autonomous Systems and Aerial Data Operations",
            "company": "ZeusVu",
            "location": "Pittsburgh, PA",
            "dates": "2018 - 2022",
            "lead": "Led FAA-compliant aerial data operations and autonomous systems concepts.",
            "bullets": [
                "Maintained a 100% safety record with no FAA incidents in complex operating environments.",
                "Developed TensorFlow machine vision concepts for inventory and field intelligence.",
                "Connected drone operations, field data, photogrammetry, restricted-airspace planning, and autonomous delivery concepts.",
            ],
        },
        {
            "title": "Operations Management / Site Leader",
            "company": "Amazon",
            "location": "Pittsburgh, PA",
            "dates": "2014 - 2018",
            "lead": "Launched Amazon Prime Pittsburgh, supporting roughly 5M second-day orders annually.",
            "bullets": [
                "Achieved 99.98% on-time delivery through cadence, escalation, staffing models, and daily performance discipline.",
                "Led and developed 1,000+ Amazon employees through launch, peak volume, daily execution, and performance cadence.",
                "Built launch rhythms for staffing, safety, throughput, quality, escalation, and customer-promise execution.",
                "Worked across robotics, Prime Air, Amazon Flex, Prime Day, proof-of-delivery, and sortation practices.",
            ],
        },
        {
            "title": "Director of Operations and Engineering Management",
            "company": "Compunetics",
            "location": "Pittsburgh, PA",
            "dates": "2000 - 2013",
            "lead": "Led engineering, production, quality, and process systems in high-reliability electronics manufacturing.",
            "bullets": [
                "Built operating systems for aerospace, defense, medical, semiconductor, and high-reliability electronics markets.",
                "Led cross-functional technical teams where engineering judgment, quality discipline, and production reality had to align.",
                "Applied Lean Six Sigma, root-cause analysis, AS9001, MIL-PRF-31032, SOPs, and technical customer translation.",
            ],
        },
    ],
    "education": [
        "Bachelor of Science, University of Pittsburgh",
        "Business Entrepreneurship, Pennsylvania University",
    ],
    "certifications": [
        "IBM Enterprise Design Thinking Practitioner",
        "Google AI Essentials",
        "Six Sigma Certification",
        "OSHA 10",
    ],
    "skills": [
        "Division IT Strategy",
        "AI Transformation",
        "ERP/Odoo Workflows",
        "SAP S/4HANA Transformation Governance",
        "Manufacturing Operations",
        "Data Governance",
        "Cyber Resilience Operating Cadence",
        "M&A Integration Readiness",
        "Workflow Automation",
        "Jira",
        "Knowledge Management",
        "KPI Design",
        "Executive Facilitation",
        "Customer Support Intelligence",
        "Robotics",
        "Supply Chain",
        "Lean Six Sigma",
        "Quality Systems",
        "Python",
        "SQL",
    ],
    "keywords": [
        "Division Vice President IT",
        "P&AI",
        "ERP Transformation",
        "SAP S/4HANA",
        "Manufacturing IT",
        "Federated IT",
        "Cybersecurity Governance",
        "Data and Analytics",
        "AI Enablement",
        "Acquisition Integration",
        "Operational Excellence",
        "Portfolio Governance",
    ],
}

COVER_LETTER_PARAGRAPHS = [
    (
        "AMETEK's P&AI Division Vice President of Information Technology role reads like a familiar operating problem: a global "
        "technical portfolio needs the benefits of standardization without losing the speed and judgment of its business units. "
        "ERP transformation, acquisition integration, cybersecurity, and data/AI adoption all have to become business outcomes, "
        "not parallel technology programs."
    ),
    (
        "That is the kind of work I do best. My background sits at the intersection of manufacturing operations, technical systems, "
        "customer promise, data visibility, and practical AI adoption. I have led teams through launch pressure, built operating "
        "cadence across production and support, and turned scattered process knowledge into workflows, dashboards, escalation paths, "
        "and training systems people can actually use."
    ),
    (
        "At Vape-Jet, I lead operations for a software-enabled hardware business where ERP/Odoo workflows, Jira, production, quality, "
        "purchasing, support, engineering issue flow, documentation, and leadership reporting all touch the same customer reality. "
        "The lesson is simple and not always comfortable: if the process owner, system truth, and daily management rhythm are unclear, "
        "technology amplifies confusion. If they are clear, technology becomes leverage."
    ),
    (
        "Amazon gave me the execution foundation. I helped launch Pittsburgh Prime operations supporting roughly five million second-day "
        "orders annually and achieved 99.98% on-time delivery through daily cadence, escalation discipline, staffing models, and customer "
        "promise management. Compunetics gave me the manufacturing foundation, where quality, engineering requirements, production reality, "
        "and customer commitments had to align in high-reliability electronics markets."
    ),
    (
        "For AMETEK, I would start by listening for the friction leaders already feel: where ERP variance is useful local fit versus hidden "
        "cost, where acquisition integrations are being reinvented, where cybersecurity asks compete with plant realities, and where AI pilots "
        "need governance before they need more enthusiasm. My working brief for the hiring team is here: "
        f"{SITE_URL}"
    ),
    (
        "My contribution would be a calm, business-first IT leadership style: protect autonomy where it creates advantage, standardize the "
        "spine where variance creates risk, and make data visible enough that leaders can run the portfolio differently. I would welcome a "
        "conversation about how that approach could help P&AI execute the next chapter."
    ),
]


def set_run_font(run, *, name="Arial", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def paragraph_border_bottom(paragraph, color="D8DEE6", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def paragraph_border_top_bottom(paragraph, color="D8DEE6", size="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for side in ("top", "bottom"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "3")
        border.set(qn("w:color"), color)
        p_bdr.append(border)


def add_para(doc, text="", *, size=8.5, after=1.6, before=0, bold=False, color=INK_RGB, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text):
    p = add_para(doc, text.upper(), size=8.7, after=2.8, before=6.2, bold=True, color=NAVY_RGB)
    paragraph_border_bottom(p)
    return p


def add_bullet(doc, text, *, size=7.55, after=1.0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, color=INK_RGB)
    return p


def add_inline_line(doc, label, items, *, size=7.15, after=0.35):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    label_run = p.add_run(f"{label} | ")
    set_run_font(label_run, size=size, bold=True, color=ACCENT_RGB)
    run = p.add_run(" | ".join(dict.fromkeys(items)))
    set_run_font(run, size=size, color=INK_RGB)
    return p


def add_scale_line(doc, text, *, size=5.72):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.4)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    paragraph_border_top_bottom(p)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=ACCENT_RGB)
    return p


def set_compact_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.34)
    section.left_margin = Inches(0.56)
    section.right_margin = Inches(0.56)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.18)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.4)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")


def build_resume_docx():
    path = OUT / f"{RESUME['filename']}.docx"
    doc = Document()
    set_compact_page(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    set_run_font(title.add_run(NAME), size=20.6, bold=True, color=NAVY_RGB)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline.paragraph_format.space_after = Pt(2)
    set_run_font(headline.add_run(RESUME["headline"]), size=8.6, bold=True, color=ACCENT_RGB)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(5)
    set_run_font(contact.add_run(CONTACT), size=7.4, color=MUTED_RGB)

    add_heading(doc, "Professional Summary")
    add_para(doc, RESUME["summary"], size=8.25, after=2.1)
    add_scale_line(doc, RESUME["scale"], size=5.72)

    add_heading(doc, RESUME["top_section_title"])
    for item in RESUME["top_bullets"]:
        add_bullet(doc, item, size=7.45, after=0.86)

    add_heading(doc, "Professional Experience")
    for role in RESUME["roles"]:
        add_para(
            doc,
            f"{role['title']} | {role['company']} | {role['location']} | {role['dates']}",
            size=8.05,
            bold=True,
            color=INK_RGB,
            after=0.45,
            before=2.25,
        )
        add_para(doc, role["lead"], size=7.55, color=MUTED_RGB, after=0.85)
        for bullet in role["bullets"]:
            add_bullet(doc, bullet, size=7.06, after=0.62)

    add_heading(doc, "Education, Certifications, and Technical Skills")
    add_inline_line(doc, "Education", RESUME["education"])
    add_inline_line(doc, "Certifications", RESUME["certifications"])
    add_inline_line(doc, "Skills", RESUME["skills"], size=6.78)
    add_inline_line(doc, "ATS Keywords", RESUME["keywords"], size=6.7, after=0)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run(f"{NAME} | {RESUME['headline']}"), size=6.8, color=RGBColor(120, 130, 138))

    doc.save(path)
    return path


def build_letter_docx():
    path = OUT / "Russell_Dudek_AMETEK_PAI_Division_VP_IT_Cover_Letter.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.56)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.0)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    set_run_font(title.add_run(NAME), size=22, bold=True, color=NAVY_RGB)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline.paragraph_format.space_after = Pt(2)
    set_run_font(headline.add_run(RESUME["headline"]), size=8.2, bold=True, color=ACCENT_RGB)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(8)
    set_run_font(contact.add_run(CONTACT), size=7.8, color=MUTED_RGB)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(9)
    paragraph_border_bottom(rule, size="8")

    add_para(doc, "July 9, 2026", size=9.3, after=5.5, color=MUTED_RGB)
    add_para(doc, "AMETEK Hiring Team", size=9.3, after=1, bold=True)
    add_para(doc, "Re: Division Vice President, Information Technology, P&AI Division", size=9.45, after=7, bold=True, color=NAVY_RGB)
    add_para(doc, "Dear AMETEK Hiring Team,", size=9.3, after=6.5)

    for paragraph in COVER_LETTER_PARAGRAPHS:
        add_para(doc, paragraph, size=9.0, after=4.8)

    add_para(doc, "Sincerely,", size=9.2, after=10)
    add_para(doc, NAME, size=9.65, after=0, bold=True, color=NAVY_RGB)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run(f"{NAME} | Cover Letter | AMETEK P&AI Division VP IT"), size=7.5, color=RGBColor(120, 130, 138))

    doc.save(path)
    return path


def pdf_footer(canvas, doc, label):
    canvas.saveState()
    canvas.setStrokeColor(RULE)
    canvas.setLineWidth(0.5)
    canvas.line(0.58 * inch, 0.38 * inch, 7.92 * inch, 0.38 * inch)
    canvas.setFillColor(MUTED)
    canvas.setFont("Helvetica", 6.8)
    canvas.drawCentredString(4.25 * inch, 0.22 * inch, f"{NAME} | {label} | Page {doc.page}")
    canvas.restoreState()


def build_resume_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("Name", fontName="Helvetica-Bold", fontSize=20.2, leading=21.6, textColor=NAVY, alignment=TA_CENTER, spaceAfter=2))
    styles.add(ParagraphStyle("Headline", fontName="Helvetica-Bold", fontSize=8.1, leading=9.8, textColor=ACCENT, alignment=TA_CENTER, spaceAfter=2))
    styles.add(ParagraphStyle("Contact", fontName="Helvetica", fontSize=7.0, leading=8.4, textColor=MUTED, alignment=TA_CENTER, spaceAfter=5))
    styles.add(ParagraphStyle("Section", fontName="Helvetica-Bold", fontSize=8.6, leading=10, textColor=NAVY, spaceBefore=6.0, spaceAfter=3, borderColor=RULE, borderWidth=0.5, borderPadding=1.8))
    styles.add(ParagraphStyle("Body", fontName="Helvetica", fontSize=7.8, leading=9.25, textColor=INK, spaceAfter=2.0))
    styles.add(ParagraphStyle("Scale", fontName="Helvetica-Bold", fontSize=7.2, leading=8.5, textColor=ACCENT, spaceAfter=1.2))
    styles.add(ParagraphStyle("Role", fontName="Helvetica-Bold", fontSize=7.8, leading=9.05, textColor=INK, spaceBefore=3.2, spaceAfter=0.8))
    styles.add(ParagraphStyle("Lead", fontName="Helvetica", fontSize=7.24, leading=8.55, textColor=MUTED, spaceAfter=0.8))
    styles.add(ParagraphStyle("ResumeBullet", fontName="Helvetica", fontSize=6.86, leading=7.86, textColor=INK, leftIndent=10, firstLineIndent=-6, spaceAfter=0.52))
    styles.add(ParagraphStyle("Small", fontName="Helvetica", fontSize=6.75, leading=7.75, textColor=INK, spaceAfter=0.55))
    return styles


def add_pdf_section(story, styles, title):
    story.append(Paragraph(title.upper(), styles["Section"]))


def add_pdf_bullets(story, styles, bullets):
    for bullet in bullets:
        story.append(Paragraph(f"- {bullet}", styles["ResumeBullet"]))


def build_resume_pdf():
    path = OUT / f"{RESUME['filename']}.pdf"
    styles = build_resume_styles()
    doc = BaseDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.58 * inch,
        rightMargin=0.58 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.48 * inch,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=lambda canvas, d: pdf_footer(canvas, d, RESUME["headline"]))])

    story = [
        Paragraph(NAME, styles["Name"]),
        Paragraph(RESUME["headline"], styles["Headline"]),
        Paragraph(CONTACT, styles["Contact"]),
    ]
    add_pdf_section(story, styles, "Professional Summary")
    story.append(Paragraph(RESUME["summary"], styles["Body"]))
    story.append(Paragraph(RESUME["scale"], styles["Scale"]))
    add_pdf_section(story, styles, RESUME["top_section_title"])
    add_pdf_bullets(story, styles, RESUME["top_bullets"])
    add_pdf_section(story, styles, "Professional Experience")
    for role in RESUME["roles"]:
        block = [
            Paragraph(f"{role['title']} | {role['company']} | {role['location']} | {role['dates']}", styles["Role"]),
            Paragraph(role["lead"], styles["Lead"]),
        ]
        for bullet in role["bullets"]:
            block.append(Paragraph(f"- {bullet}", styles["ResumeBullet"]))
        story.append(KeepTogether(block))
    add_pdf_section(story, styles, "Education, Certifications, and Technical Skills")
    story.append(Paragraph("<b>Education</b> | " + " | ".join(RESUME["education"]), styles["Small"]))
    story.append(Paragraph("<b>Certifications</b> | " + " | ".join(RESUME["certifications"]), styles["Small"]))
    story.append(Paragraph("<b>Skills</b> | " + " | ".join(dict.fromkeys(RESUME["skills"])), styles["Small"]))
    story.append(Paragraph("<b>ATS Keywords</b> | " + " | ".join(dict.fromkeys(RESUME["keywords"])), styles["Small"]))
    doc.build(story)
    return path


def build_letter_pdf():
    path = OUT / "Russell_Dudek_AMETEK_PAI_Division_VP_IT_Cover_Letter.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("Name", fontName="Helvetica-Bold", fontSize=22, leading=24, textColor=NAVY, alignment=TA_CENTER, spaceAfter=2))
    styles.add(ParagraphStyle("Headline", fontName="Helvetica-Bold", fontSize=8.0, leading=9.6, textColor=ACCENT, alignment=TA_CENTER, spaceAfter=2))
    styles.add(ParagraphStyle("Contact", fontName="Helvetica", fontSize=7.3, leading=8.8, textColor=MUTED, alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle("Body", fontName="Helvetica", fontSize=9.0, leading=11.25, textColor=INK, spaceAfter=6.0))
    styles.add(ParagraphStyle("Meta", fontName="Helvetica", fontSize=9.0, leading=11.0, textColor=MUTED, spaceAfter=6.0))
    styles.add(ParagraphStyle("Subject", fontName="Helvetica-Bold", fontSize=9.2, leading=11.2, textColor=NAVY, spaceAfter=8.0))
    styles.add(ParagraphStyle("Sig", fontName="Helvetica-Bold", fontSize=9.5, leading=11.2, textColor=NAVY, spaceAfter=0))
    doc = BaseDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.72 * inch,
        rightMargin=0.72 * inch,
        topMargin=0.58 * inch,
        bottomMargin=0.58 * inch,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=lambda canvas, d: pdf_footer(canvas, d, "Cover Letter | AMETEK P&AI Division VP IT"))])

    story = [
        Paragraph(NAME, styles["Name"]),
        Paragraph(RESUME["headline"], styles["Headline"]),
        Paragraph(CONTACT, styles["Contact"]),
        Paragraph("July 9, 2026", styles["Meta"]),
        Paragraph("AMETEK Hiring Team", styles["Body"]),
        Paragraph("Re: Division Vice President, Information Technology, P&AI Division", styles["Subject"]),
        Paragraph("Dear AMETEK Hiring Team,", styles["Body"]),
    ]
    for paragraph in COVER_LETTER_PARAGRAPHS:
        story.append(Paragraph(paragraph, styles["Body"]))
    story.extend([Spacer(1, 6), Paragraph("Sincerely,", styles["Body"]), Paragraph(NAME, styles["Sig"])])
    doc.build(story)
    return path


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    paths = [
        build_resume_docx(),
        build_resume_pdf(),
        build_letter_docx(),
        build_letter_pdf(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
